import re
from docx import Document
from datetime import datetime, timedelta
from .models import DocxModel


def check(user_filter):
    """Функция для получени файла и запуска скрипта"""
    query = DocxModel.objects.filter(owner__user__username=user_filter).last()
    with open(query.document.path, 'rb') as f:
        doc = Document(f).tables[0]
        return func(doc, query)


def has_cyrillic(text):
    return bool(re.search('[а-яА-Я]', text))


def check_work_shift(query):
    """Функция для определения текущей смены работника"""
    if query.work_shift == 'д':
        data = datetime.today()
        return data.strftime('%d.%m.%Y')
    else:
        data = datetime.today() + timedelta(days=1)
        return data.strftime('%d.%m.%Y')


def func(table, query):
    """Функция для работы с основной таблицей файла"""
    list_full = []
    list_short = []
    list_end_first = []
    list_end_last = []
    count = 0
    try:
        for row in table.rows[2:]:
            count += 1
            string = ''
            for cell in row.cells[1:3]:
                i = cell.text.rstrip('\n ')
                string = f"{string + i + ', '}"
            new_line = f"-\t{string[0].lower() + string[1:]}".rstrip(', ')
            list_full.append(new_line)
    except IndexError as e:
        print(e)
        return 'Что-то пошло не так с первой частью документа, а именно со 2 по 3 колонку'

    try:
        for row in table.rows[2:]:
            cell = row.cells[1].text.rstrip('\n')
            new_string = f"-\t{cell[0].lower() + cell[1:]}"
            list_short.append(new_string)
    except IndexError as e:
        print(e)
        return 'Что-то пошло не так со второй частью документа, а именно со 2 колонкой'

    try:
        for row in table.rows[2:]:
            list_end_first.append(row.cells[5].text)
    except IndexError as e:
        print(e)
        return 'Что-то пошло не так при парсинге документа, а именно с 6 колонкой'
    try:
        for row in table.rows[2:]:
            context = row.cells[6].text
            if context != '':
                second_part = context.split('\n')
                list_end_last.append(second_part[1])
            else:
                list_end_last.append(context)
    except IndexError as e:
        print(e)
        return 'Что-то пошло не так при парсинге документа, а именно с 7 колонкой'

    writer(list_full, list_short, list_end_first, list_end_last, count, query)


def writer(list_fl, list_st, list_end_ft, list_end_st, ct, query):
    """Функция для записи данных в файл"""
    with open(query.document.path, 'rb') as f:
        doc = Document(f)
    """
    Два генератора списков для удаления всех данных из файла и записи в чистый документ
    """
    [par._element.getparent().remove(par._element) for par in doc.paragraphs]
    [tab._element.getparent().remove(tab._element) for tab in doc.tables]

    date_time = check_work_shift(query)
    doc.add_paragraph(f'Выявлено {ct} нарушений\n')

    for par_fl in range(len(list_fl)):  # Цикл для формирования первой части документа
        if has_cyrillic(list_end_ft[par_fl]):
            content = f'\n{list_fl[par_fl]}. '
            doc.add_paragraph(content).add_run(f'{list_end_ft[par_fl]}\n').bold = True
        else:
            content = f'\n{list_fl[par_fl]}\n'
            doc.add_paragraph(content)

    doc.add_paragraph(f'Выдано предписание № {query.number_check}-1 от '
                      f'{date_time}\n').runs[0].bold = True

    for par_st in range(len(list_st[:-1])):  # Цикл для формирования второй части документа
        if has_cyrillic(list_end_ft[par_st]):
            content = list_st[par_st]
            if list_end_st[par_st] != '':
                six_column = doc.add_paragraph(content).add_run(f'. {list_end_ft[par_st]}')
                six_column.add_run(f'. {list_end_st[par_st]};')

            else:
                doc.add_paragraph(content).add_run(f'. {list_end_ft[par_st]};')

        else:
            if list_end_st[par_st] != '':
                content = list_st[par_st]
                doc.add_paragraph(content).add_run(f'. {list_end_st[par_st]};')
            else:
                content = f'{list_st[par_st]};'
                doc.add_paragraph(content)

    if has_cyrillic(list_end_ft[-1]):
        content = list_st[-1]
        if list_end_st[-1] != '':
            six_column = doc.add_paragraph(content).add_run(f'. {list_end_ft[-1]}')
            six_column.add_run(f'. {list_end_st[-1]}.')
        else:
            doc.add_paragraph(content).add_run(f'. {list_end_ft[-1]}.')
    else:
        if list_end_st[-1] != '':
            content = list_st[-1]
            doc.add_paragraph(content).add_run(f'. {list_end_st[-1]}.')
        else:
            content = f"{list_st[-1]}."
            doc.add_paragraph(content)

    doc.save(query.document.path)

